import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.text import WD_BREAK
import re
import io

# ==============================================================================
# 1. 核心逻辑与算法模块 (Backend Logic)
# ==============================================================================

def set_global_document_settings(doc):
    """
    全局设置：页边距 (1英寸)
    注意：APA 7th 要求所有页边距均为 1 英寸 (2.54cm)
    """
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def apply_basic_font_style(paragraph):
    """
    基础样式应用：Times New Roman, 12pt, 双倍行距
    """
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE # 双倍行距
    
    # 有些文档可能混杂了复杂的 Style，这里强制覆盖 Run 级别的字体
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    # 为了保险，也尝试设置 Style 级别（如果有 Normal 样式）
    try:
        style = paragraph.style
        if style and hasattr(style, 'font'):
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
    except:
        pass

def locate_structural_indices(doc, has_title_page):
    """
    智能定位算法：
    1. 寻找 body_start_index (正文开始的段落索引)
    2. 寻找 ref_start_index (参考文献开始的段落索引)
    """
    paragraphs = doc.paragraphs
    total_pars = len(paragraphs)
    
    body_start_index = 0
    ref_start_index = total_pars # 默认为末尾，即没找到

    # --- A. 定位参考文献 (Reference) ---
    # 策略：倒序或正序查找 "Reference" 独占一行的段落
    # 优先找 References，这样可以确定正文的边界
    for i, p in enumerate(paragraphs):
        # 清洗文本：去空格，转小写
        text = p.text.strip().lower()
        # 匹配 "reference" 或 "references"，且字数不能太多（防止匹配到正文里的句子）
        if text in ['reference', 'references'] or text == 'reference list':
            ref_start_index = i
            break
    
    # --- B. 定位正文起始 (Body Start) ---
    # 这是一个高难度动作，涉及“安全视窗”和“非空穿透”逻辑
    
    if has_title_page:
        found_page_break = False
        SAFE_SEARCH_LIMIT = 50 # 安全视窗：只在前50段寻找标题页逻辑
        search_limit = min(SAFE_SEARCH_LIMIT, ref_start_index)
        
        # 策略 1: 寻找物理分页符 (Hard Page Break)
        for i in range(search_limit):
            # 深入 XML 检查是否有 <w:br w:type="page"/>
            if '<w:br w:type="page"/>' in p._element.xml:
                body_start_index = i + 1 # 分页符所在段落的下一段是正文
                found_page_break = True
                break
        
        # 策略 2: 软换行穿透 (Rule of 6)
        if not found_page_break:
            non_empty_count = 0
            target_index = 0
            
            # 1. 计数：找到第6个有文字的段落 (通常是 Date)
            for i in range(search_limit):
                if paragraphs[i].text.strip():
                    non_empty_count += 1
                if non_empty_count == 6:
                    target_index = i
                    break
            
            # 2. 穿透：从第6个非空段落往后，跳过所有空行，直到遇到文字
            for j in range(target_index + 1, search_limit):
                if paragraphs[j].text.strip():
                    body_start_index = j
                    break
                    
    return body_start_index, ref_start_index

def delete_paragraph(paragraph):
    """
    辅助函数：彻底删除一个段落对象
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def process_formatting(doc, config):
    """
    主处理逻辑 (V3.1 Final: Pixel-Perfect Title Page)
    """
    # 1. 全局设置
    set_global_document_settings(doc)
    
    # 2. 定位结构
    body_start, ref_start = locate_structural_indices(doc, config['has_title_page'])
    
    paragraphs = doc.paragraphs
    
    # ==========================
    # 阶段 0: 标题页特殊处理 (Title Page Formatting)
    # ==========================
    if config['has_title_page'] and body_start > 0:
        title_lines_count = 0
        last_title_paragraph = None
        
        # 先把正文第一段的对象存下来，因为后面删行会导致索引变化，但对象引用不变
        # 这一步非常关键，用于最后在其前方插入分页符
        first_body_paragraph = paragraphs[body_start]
        
        # A. 格式化标题页的所有段落 (包括空行)
        for i in range(body_start):
            p = paragraphs[i]
            
            # 修正 1: 无论是不是空行，都强制应用双倍行距和字体
            # 这样标题上方的空行高度才会正确，标题位置才不会偏上
            apply_basic_font_style(p)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 如果是有字的行，进行计数和加粗处理
            if p.text.strip():
                title_lines_count += 1
                last_title_paragraph = p 
                
                # 第一行加粗 (文章主标题)
                if title_lines_count == 1:
                    for run in p.runs:
                        run.bold = True

        # B. 清洗标题页与正文之间的“垃圾空行”
        # 策略：找到 last_title_paragraph 的索引，删除它之后直到 body_start 之间的所有段落
        if last_title_paragraph:
            # 重新定位 last_title_paragraph 的索引
            last_title_idx = -1
            for idx in range(body_start):
                if paragraphs[idx] == last_title_paragraph:
                    last_title_idx = idx
                    break
            
            # 倒序删除中间的空行
            if last_title_idx != -1:
                for idx in range(body_start - 1, last_title_idx, -1):
                    # 只删空行 (防止误删内容)
                    if not paragraphs[idx].text.strip():
                        delete_paragraph(paragraphs[idx])
            
            # C. 修正 2: 在新的一行插入分页符 (Aesthetic Page Break)
            # 逻辑：在 first_body_paragraph (正文第一段) 的前面，插入一个新的空白段落
            # 然后在这个新段落里放分页符。这样分页符就独占一行，不会挤在日期后面了。
            
            # 检查是否原本就有分页符 (避免双重分页)
            has_existing_break = False
            # 检查 last_title_paragraph 里面有没有
            if '<w:br w:type="page"/>' in last_title_paragraph._element.xml:
                has_existing_break = True
            # 检查 first_body_paragraph 里面有没有 (有时候分页符在正文开头)
            if '<w:br w:type="page"/>' in first_body_paragraph._element.xml:
                has_existing_break = True
            
            if not has_existing_break:
                # 插入这一行“缓冲带”
                spacer_p = first_body_paragraph.insert_paragraph_before()
                # 给这个分页符段落也加上标准格式 (虽然看不见，但为了规范)
                apply_basic_font_style(spacer_p) 
                # 添加分页符
                spacer_p.add_run().add_break(WD_BREAK.PAGE)

    # ==========================
    # 阶段 I: 处理正文 (Body)
    # ==========================
    # 刷新 paragraphs 列表 (因为刚刚删了行，又插了行)
    paragraphs = doc.paragraphs
    
    # 重新定位 body_start
    # 简单粗暴且有效的方法：重新跑一次定位，或者直接找 Page Break 的位置
    # 由于我们刚刚确保插入了 Page Break，现在找 Page Break 是最稳的
    
    new_body_start = 0
    if config['has_title_page']:
        for i, p in enumerate(paragraphs):
            # 寻找刚刚插入的那个分页符段落，正文在它下一行
            if '<w:br w:type="page"/>' in p._element.xml:
                new_body_start = i + 1
                break
    
    # 重新定位 ref_start (因为行数变了)
    _, new_ref_start = locate_structural_indices(doc, False)
    
    # 开始处理正文
    for i in range(new_body_start, new_ref_start):
        p = paragraphs[i]
        text = p.text.strip()
        
        if not text: continue
            
        apply_basic_font_style(p)
        pf = p.paragraph_format
        
        # Case 1: 文章主标题
        if i == new_body_start and config['has_article_title']:
            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pf.first_line_indent = Inches(0)
            for run in p.runs:
                run.bold = True
                
        # Case 2: 二级标题
        elif len(text.split()) < 15 and text[-1] not in ['.', ':', '?', '!']:
            pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            pf.first_line_indent = Inches(0)
            pf.left_indent = Inches(0)
            for run in p.runs:
                run.bold = True
                
        # Case 3: 普通正文
        else:
            pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            pf.first_line_indent = Inches(0.5)
            pf.left_indent = Inches(0) 

    # ==========================
    # 阶段 II: 处理参考文献 (Refs)
    # ==========================
    if new_ref_start < len(paragraphs):
        # 强制分页逻辑
        if new_ref_start > 0:
            prev_p_ref = paragraphs[new_ref_start - 1]
            if '<w:br w:type="page"/>' not in prev_p_ref._element.xml:
                 prev_p_ref.add_run().add_break(WD_BREAK.PAGE)

        ref_title_p = paragraphs[new_ref_start]
        ref_title_p.text = "References" 
        apply_basic_font_style(ref_title_p)
        ref_title_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ref_title_p.paragraph_format.first_line_indent = Inches(0)
        for run in ref_title_p.runs:
            run.bold = True
            
        ref_entries = []
        for i in range(new_ref_start + 1, len(paragraphs)):
            p = paragraphs[i]
            if p.text.strip():
                ref_entries.append(p.text.strip())

        # 重写 Reference 逻辑
        for i in range(len(paragraphs) - 1, new_ref_start, -1):
            delete_paragraph(paragraphs[i])
            
        if config['sort_references']:
            ref_entries.sort()
            
        for entry in ref_entries:
            new_p = doc.add_paragraph(entry)
            apply_basic_font_style(new_p)
            new_p.paragraph_format.first_line_indent = Inches(-0.5)
            new_p.paragraph_format.left_indent = Inches(0.5)

    return doc

def check_missing_citations(doc):
    """
    双向引用检查逻辑 (V4 Updated: Two-way & Author+Year Key)
    """
    import re
    
    # 1. 获取全文文本
    # 为了避免匹配到页眉页脚或 Reference 列表本身，我们需要界定范围
    # 简单起见，我们假设全文文本就是 process_formatting 之后的 doc 对象
    # 但为了精准，我们只提取 Ref 标题之前的内容作为 "Body Text"
    paragraphs = doc.paragraphs
    body_text = ""
    ref_text_list = []
    
    found_ref_section = False
    for p in paragraphs:
        txt = p.text.strip()
        # 简单的状态机，找到 References 标题后切换状态
        if txt.lower() == 'references' or txt.lower() == 'reference list':
            found_ref_section = True
            continue
            
        if not found_ref_section:
            body_text += txt + " "
        else:
            if txt:
                ref_text_list.append(txt)

    # ==========================================
    # Step A: 解析参考文献列表 (Reference List)
    # 目标：提取 (First_Author_Lastname, Year)
    # ==========================================
    ref_keys = set()
    ref_details = {} # 用于存储原始文本，方便展示
    
    for ref_item in ref_text_list:
        # 策略：
        # 1. 年份：找圆括号里的4位数字，通常在行首附近
        #    Regex: 匹配行首开始的任意字符，直到发现 (20xx) 或 (n.d.)
        # 2. 作者：年份之前的部分，取第一个单词作为姓氏
        
        # 匹配年份：(2019) 或 (n.d.)
        year_match = re.search(r'\((\d{4}|n\.d\.)\)', ref_item)
        
        if year_match:
            year = year_match.group(1)
            
            # 提取作者：取年份括号之前的所有文本
            pre_year_text = ref_item[:year_match.start()]
            
            # 提取第一个单词作为姓氏 (移除逗号等标点)
            # 比如 "Wang, I. (2020)" -> "Wang"
            # 比如 "World Health Organization (2020)" -> "World" (虽然不完美，但够用)
            if pre_year_text:
                # 简单的 split 逻辑
                first_author = pre_year_text.split(',')[0].strip().split(' ')[0]
                
                # 清洗一下非字母字符 (比如有些 Ref 前面有奇怪的编号)
                first_author = re.sub(r'[^a-zA-Z\u4e00-\u9fa5]', '', first_author)
                
                if first_author and year:
                    key = (first_author.lower(), year)
                    ref_keys.add(key)
                    # 存储一下原始文本供报告使用
                    ref_details[key] = ref_item[:50] + "..." # 只存前50个字符

    # ==========================================
    # Step B: 解析正文引用 (In-text Citations)
    # 目标：提取 (Author, Year)
    # ==========================================
    body_keys = set()
    
    # Regex 策略：
    # 1. 寻找括号内容 (...)
    # 2. 括号内必须包含年份 \d{4}
    # 3. 排除 (see Table 1) 这种非引用
    
    # 匹配所有括号内容
    parentheses_content = re.findall(r'\(([^)]+)\)', body_text)
    
    for content in parentheses_content:
        # 1. 必须包含年份 (20xx) 或 n.d.
        if not re.search(r'\d{4}|n\.d\.', content):
            continue
            
        # 2. 可能包含多个引用，用分号 ; 隔开
        # 例如: (Wang, 2020; Zhang & Li, 2021)
        citations = content.split(';')
        
        for cite in citations:
            cite = cite.strip()
            
            # 再次确认这一小段里有年份
            year_match = re.search(r'(\d{4}|n\.d\.)', cite)
            if not year_match:
                continue
                
            year = year_match.group(1)
            
            # 提取作者部分：年份前面的文本
            # 比如 "Wang et al., 2020" -> "Wang et al.,"
            # 比如 "Zhang & Li, 2021" -> "Zhang & Li,"
            author_part = cite[:year_match.start()].strip()
            
            # 过滤干扰词 (e.g., see, cf.)
            ignore_words = ['see', 'e.g.', 'cf.', 'also', 'table', 'figure']
            is_ignored = False
            for word in ignore_words:
                if word in author_part.lower():
                    # 如果包含干扰词，尝试清洗，取干扰词之后的部分
                    # 简单处理：如果整个部分就是干扰词（如 (Figure 1)），由于前面校验了年份，这里很难误判
                    # 但如果是 (see Wang, 2020)，我们需要去掉 "see"
                    pass 
            
            # 提取姓氏：取第一个单词
            # 处理 "Wang et al." -> Wang
            # 处理 "Zhang & Li" -> Zhang
            tokens = re.split(r'[\s,&]+', author_part) # 按空格、逗号、& 分割
            
            clean_tokens = [t for t in tokens if t and t.lower() not in ignore_words]
            
            if clean_tokens:
                first_author = clean_tokens[0]
                # 清洗
                first_author = re.sub(r'[^a-zA-Z\u4e00-\u9fa5]', '', first_author)
                
                if first_author:
                    body_keys.add((first_author.lower(), year))

    # ==========================================
    # Step C: 双向对比 (Two-way Match)
    # ==========================================
    
    # 1. 正文有，Ref 列表没有 (Missing in Refs)
    missing_in_refs = body_keys - ref_keys
    
    # 2. Ref 列表有，正文没有 (Missing in Body)
    missing_in_body = ref_keys - body_keys
    
    return list(missing_in_refs), list(missing_in_body)

# ==============================================================================
# 2. 前端交互模块 (Frontend UI)
# ==============================================================================

def main():
    st.set_page_config(page_title="APA 7th Format Helper", page_icon="🎓")

# --- CSS 注入：美化 & 修复侧边栏 & 高亮复制按钮 ---
    hide_streamlit_style = """
                <style>
                /* 1. 找回侧边栏和菜单 */
                #MainMenu {visibility: visible;} 
                
                /* 2. 隐藏页脚 */
                footer {visibility: hidden;}
                
                /* 3. 底部签名 */
                .custom-footer {
                    position: fixed;
                    left: 0;
                    bottom: 0;
                    width: 100%;
                    background-color: #f0f2f6;
                    color: #555;
                    text-align: center;
                    padding: 10px;
                    font-size: 14px;
                    font-family: 'Arial', sans-serif;
                    border-top: 1px solid #e6e6e6;
                    z-index: 999;
                }
                
                /* 4. 高亮检查报告的复制按钮 */
                /* 针对 Streamlit 的代码块复制按钮进行样式覆盖 */
                [data-testid="stCopyButton"] {
                    background-color: #FF4B4B !important; /* 显眼的红色背景，或者换成你喜欢的蓝色 #4B9EFF */
                    color: white !important;
                    opacity: 1 !important; /* 强制不透明 */
                    border: 1px solid white !important;
                    border-radius: 4px !important;
                    transform: scale(1.1); /* 稍微放大一点 */
                }
                
                /* 鼠标悬停时的效果 */
                [data-testid="stCopyButton"]:hover {
                    background-color: #FF2B2B !important;
                    transform: scale(1.2);
                }
                </style>
                """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)

    # --- 标题区 ---
    st.title("📄 APA 7th Format Assistant")
    # st.markdown("Designed specifically for **Dr. Jin**'s academic workflow.")
    st.markdown("---")

    # --- 侧边栏配置 ---
    st.sidebar.header("⚙️ Configuration")
    
    has_title_page = st.sidebar.checkbox(
        "Has Title Page? ", 
        value=False,
        help="勾选后，工具将智能跳过封面页（识别分页符或前6行内容），从第二页开始格式化。"
    )
    
    has_article_title = st.sidebar.checkbox(
        "Has Article Title?", 
        value=True,
        help="勾选后，正文的第一段将被格式化为居中加粗的主标题。"
    )
    
    sort_references = st.sidebar.checkbox(
        "Auto-sort References (A-Z)", 
        value=False,
        help="勾选后，参考文献列表将被自动按字母顺序排序。请注意，这可能会移除斜体格式（如期刊名）。"
    )
    
    # 动态警告
    if sort_references:
        st.sidebar.warning(
            "⚠️ Warning: Auto-sorting will verify strict alphabetical order but "
            "**MAY REMOVE ITALICS** (e.g., journal names). Uncheck if you want to keep existing italics."
        )

    check_citations_opt = st.sidebar.checkbox(
        "Check Missing Citations", 
        value=True,
        help="生成一份报告，检查正文中引用的文献是否在 Reference 列表中缺失。"
    )

    # --- 文件上传区 ---
    uploaded_file = st.file_uploader("Drop your dissertation/paper here (.docx)", type="docx")

    if uploaded_file is not None:
        try:
            # 读取文件
            doc = Document(uploaded_file)
            
            # --- 运行处理逻辑 ---
            processed_doc = process_formatting(doc, {
                'has_title_page': has_title_page,
                'has_article_title': has_article_title,
                'sort_references': sort_references
            })
            
            st.success("✅ Formatting complete! Ready for download.")
            
# --- 引用检查报告 (V4 Updated: Two-way Report) ---
            if check_citations_opt:
                # 解包返回的两个列表
                missing_in_refs, missing_in_body = check_missing_citations(doc)
                
                report_content = ""
                
                # 头部提示
                if sort_references:
                    report_content += "⚠️ [WARNING] Auto-sort is ON. Italics in References removed.\n"
                else:
                    report_content += "ℹ️ [INFO] Auto-sort is OFF. Formatting checks only.\n"
                
                report_content += "-" * 40 + "\n"
                
                # --- Part 1: 正文引了，文献表没列 (最严重) ---
                if missing_in_refs:
                    report_content += "🚨 CITED IN TEXT BUT MISSING IN REFERENCES:\n"
                    report_content += "(Please verify spelling or year matches)\n\n"
                    for author, year in missing_in_refs:
                        # 把名字首字母大写，看起来更正规
                        report_content += f"[ ] {author.title()}, {year}\n"
                else:
                    report_content += "✅ All in-text citations found in Reference list.\n"
                
                report_content += "\n" + "-" * 40 + "\n"
                
                # --- Part 2: 文献表列了，正文没引 (冗余) ---
                if missing_in_body:
                    report_content += "❓ LISTED IN REFERENCES BUT NOT FOUND IN TEXT:\n"
                    report_content += "(Did you forget to cite these?)\n\n"
                    for author, year in missing_in_body:
                        report_content += f"[ ] {author.title()}, {year}\n"
                else:
                    report_content += "✅ All references are cited in the text.\n"

                report_content += "\n*Report generated by APA 7th Format Assistant*"

                # UI 展示
                st.warning("🧐 **Citation Check Report:**")
                st.code(report_content, language="markdown")
                st.caption("*Click the red copy button (top-right) to grab this report.*")

            # --- 导出 ---
            bio = io.BytesIO()
            processed_doc.save(bio)
            
            # 构建新文件名
            original_name = uploaded_file.name.rsplit('.', 1)[0]
            new_name = f"{original_name}_APA_Formatted.docx"
            
            st.download_button(
                label="📥 Download Formatted Document",
                data=bio.getvalue(),
                file_name=new_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error("Oops! Something went wrong processing the file.")
            st.error(f"Error details: {e}")

    # --- 底部签名 (Inject Footer) ---
    st.markdown(
        """
        <div class="custom-footer">
            Designed specially for Dr. Jin
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
